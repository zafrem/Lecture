# pip install python-docx
import os

import _1_technical_analysis_Stochastic as stochastic
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


if "__main__" == __name__:
    # Create Image
    screenshop_file_name = 'stochastic_report.png'
    report_docx = 'stochastic_report.docx'

    if os.path.exists(screenshop_file_name):
        os.remove(screenshop_file_name)

    df = stochastic.init_pre_df()
    plt, df = stochastic.stochastic(df)
    plt.savefig(screenshop_file_name, bbox_inches='tight')

    document = Document()
    # Add a title
    title = document.add_heading('Stochastic Report', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add a paragraph with bold and italic text
    paragraph = document.add_paragraph('This is a sample report created using the python-docx library.')
    run = paragraph.runs[0]
    run.bold = True
    run.italic = True

    # Add a heading
    document.add_heading('Section 1: Introduction', level=2)

    # Add a bulleted list
    list_paragraph = document.add_paragraph()
    list_paragraph.add_run('This data is BTC to USD in binance site.').bold = True
    list_paragraph.add_run('\n')
    list_paragraph.add_run("This graph was plotted with data where random numbers were added to create the maximum and minimum values for the day. It was generated as a sample to create a report, so it's not real data.")

    # Add an images
    document.add_heading('Section 2: Plot', level=2)
    document.add_paragraph('Here is Stochastic Plot')
    document.add_picture(screenshop_file_name, width=Pt(300))

    # Add a table
    document.add_heading('Section 3: Data Table', level=2)
    table = document.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(100)
    table.cell(0, 0).text = 'Date'
    table.cell(0, 1).text = 'Min'
    table.cell(0, 2).text = 'Max'
    for i, data in enumerate(
            [('2025-01-28 17:53:58', '2673.5', '2973.5'), ('2025-01-28 17:54:08', '2675.5', '2680.5'), ('2025-01-28 17:54:18', '2103.5', '2673.5')], start=1):
        table.cell(i, 0).text = data[0]
        table.cell(i, 1).text = data[1]
        table.cell(i, 2).text = data[2]

    if os.path.exists(report_docx):
        os.remove(report_docx)

    document.save(report_docx)
    os.remove(screenshop_file_name)